import spacy
import docx as docx_lib
import pdfplumber
import streamlit as st
import sqlite3
import pandas as pd
import re
import json
import uuid
import io
from datetime import datetime
from pathlib import Path

# ══════════════════════════════════════════════════════════
#  DEPENDENCY CHECK & IMPORTS
# ══════════════════════════════════════════════════════════


def check_and_import():
    missing = []
    try:
        import pdfplumber
    except ImportError:
        missing.append("pdfplumber")
    try:
        import docx
    except ImportError:
        missing.append("python-docx")
    try:
        import spacy
    except ImportError:
        missing.append("spacy")
    return missing


missing_pkgs = check_and_import()

# ══════════════════════════════════════════════════════════
#  PAGE CONFIG  (must be first Streamlit call)
# ══════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Resume Parser AI",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ══════════════════════════════════════════════════════════
#  GLOBAL CSS
# ══════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Sans:wght@300;400;500;600&display=swap');

:root {
    --bg:        #080c18;
    --surface:   #0f1424;
    --surface2:  #161d30;
    --surface3:  #1c2540;
    --accent:    #00d4aa;
    --accent2:   #6c8dff;
    --accent3:   #f59e0b;
    --pink:      #f472b6;
    --text:      #e2e8f0;
    --muted:     #64748b;
    --border:    rgba(0,212,170,0.12);
    --border2:   rgba(108,141,255,0.12);
    --radius:    14px;
}

html, body, [class*="css"] {
    background-color: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
}
[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border);
}

/* ── Hero ───────────────────────────────────────────── */
.hero {
    background: linear-gradient(135deg, var(--surface) 0%, var(--surface2) 100%);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 2.2rem 2rem 1.8rem;
    text-align: center;
    margin-bottom: 1.5rem;
    position: relative;
    overflow: hidden;
}
.hero::before {
    content:'';
    position:absolute; top:-60%; left:-40%;
    width:180%; height:180%;
    background: radial-gradient(ellipse at 40% 40%, rgba(0,212,170,0.07) 0%, transparent 55%),
                radial-gradient(ellipse at 70% 70%, rgba(108,141,255,0.06) 0%, transparent 50%);
    pointer-events:none;
}
.hero-title {
    font-family: 'Syne', sans-serif !important;
    font-size: 2.5rem; font-weight: 800;
    background: linear-gradient(135deg, #00d4aa 0%, #6c8dff 55%, #f472b6 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
    margin: 0; letter-spacing: -1px;
}
.hero-sub { color: var(--muted); font-size: 0.95rem; margin-top: 0.4rem; }

/* ── Cards ──────────────────────────────────────────── */
.card {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1.3rem 1.4rem;
    margin-bottom: 1rem;
}
.card-title {
    font-family: 'Syne', sans-serif;
    font-size: 1rem; font-weight: 700;
    color: var(--accent); margin-bottom: 0.8rem;
    display: flex; align-items: center; gap: 6px;
}

/* ── Metric cards ───────────────────────────────────── */
.metric-card {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: var(--radius);
    padding: 1.2rem; text-align: center;
}
.metric-val   { font-family:'Syne',sans-serif; font-size:2rem; font-weight:800; }
.metric-label { font-size:0.78rem; color:var(--muted); margin-top:4px; }

/* ── Info rows ──────────────────────────────────────── */
.info-row {
    display: flex; align-items: flex-start;
    gap: 10px; padding: 0.55rem 0;
    border-bottom: 1px solid rgba(255,255,255,0.04);
}
.info-row:last-child { border-bottom: none; }
.info-label {
    font-size: 0.75rem; font-weight: 600; color: var(--muted);
    text-transform: uppercase; letter-spacing: 0.8px;
    min-width: 100px; padding-top: 2px;
}
.info-val { font-size: 0.92rem; color: var(--text); flex:1; line-height:1.5; }

/* ── Skill tags ─────────────────────────────────────── */
.skill-wrap { display:flex; flex-wrap:wrap; gap:7px; margin-top:4px; }
.skill-tag {
    background: rgba(0,212,170,0.1);
    border: 1px solid rgba(0,212,170,0.25);
    color: #00d4aa; font-size: 0.78rem; font-weight: 600;
    padding: 3px 12px; border-radius: 20px;
    transition: all 0.2s;
}

/* ── Upload zone ────────────────────────────────────── */
[data-testid="stFileUploader"] {
    border: 2px dashed rgba(0,212,170,0.25) !important;
    border-radius: var(--radius) !important;
    background: rgba(0,212,170,0.03) !important;
}

/* ── Buttons ────────────────────────────────────────── */
.stButton button {
    background: linear-gradient(135deg, var(--accent), var(--accent2)) !important;
    color: #000 !important; border: none !important;
    border-radius: 10px !important; font-weight: 700 !important;
    font-family: 'DM Sans', sans-serif !important;
}

/* ── Progress bar ───────────────────────────────────── */
.parse-score {
    display: flex; align-items: center; gap: 12px;
    margin-top: 8px;
}
.score-bar-wrap {
    flex: 1; height: 8px;
    background: rgba(255,255,255,0.06);
    border-radius: 4px; overflow: hidden;
}
.score-bar-fill {
    height: 100%; border-radius: 4px;
    background: linear-gradient(90deg, #00d4aa, #6c8dff);
}
.score-pct { font-family:'Syne',sans-serif; font-weight:700; font-size:0.9rem; color:var(--accent); }

/* ── Sidebar brand ──────────────────────────────────── */
.sb-title {
    font-family:'Syne',sans-serif; font-size:1.35rem; font-weight:800;
    background:linear-gradient(135deg,#00d4aa,#6c8dff);
    -webkit-background-clip:text; -webkit-text-fill-color:transparent;
}

/* ── Section headers ────────────────────────────────── */
.sec-head {
    font-family:'Syne',sans-serif; font-size:1.3rem; font-weight:800;
    color:var(--text); margin-bottom:1rem;
}

/* ── Badge ──────────────────────────────────────────── */
.badge-new {
    background:rgba(0,212,170,0.15); color:#00d4aa;
    border:1px solid rgba(0,212,170,0.3);
    font-size:0.7rem; font-weight:700; padding:2px 9px;
    border-radius:20px; vertical-align:middle; margin-left:8px;
}

/* ── Table ──────────────────────────────────────────── */
.stDataFrame { border-radius:10px; overflow:hidden; }

::-webkit-scrollbar { width:6px; }
::-webkit-scrollbar-track { background:var(--bg); }
::-webkit-scrollbar-thumb { background:rgba(0,212,170,0.2); border-radius:3px; }
hr { border-color:rgba(0,212,170,0.1) !important; margin:1rem 0; }
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════
#  SHOW INSTALL INSTRUCTIONS IF MISSING PACKAGES
# ══════════════════════════════════════════════════════════
if missing_pkgs:
    st.error("⚠️ Missing required packages! Run the commands below in your terminal, then restart the app.")
    cmds = []
    if "pdfplumber" in missing_pkgs:
        cmds.append("pip install pdfplumber")
    if "python-docx" in missing_pkgs:
        cmds.append("pip install python-docx")
    if "spacy" in missing_pkgs:
        cmds.append("pip install spacy")
        cmds.append("python -m spacy download en_core_web_sm")
    st.code("\n".join(cmds), language="bash")
    st.stop()

# Safe imports after check

# Load spacy model


@st.cache_resource
def load_spacy():
    try:
        return spacy.load("en_core_web_sm")
    except OSError:
        st.error(
            "spaCy model not found! Run: `python -m spacy download en_core_web_sm`")
        st.stop()


nlp = load_spacy()

# ══════════════════════════════════════════════════════════
#  SKILLS DICTIONARY
# ══════════════════════════════════════════════════════════
SKILLS_DB = {
    # Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "c", "ruby", "go", "rust", "kotlin", "swift",
    "php", "scala", "r", "matlab", "perl", "bash", "shell", "dart", "lua", "haskell", "elixir", "clojure",
    # Web
    "html", "css", "react", "angular", "vue", "next.js", "nuxt", "svelte", "jquery", "bootstrap", "tailwind",
    "sass", "less", "webpack", "vite", "node.js", "express", "django", "flask", "fastapi", "spring", "rails",
    "asp.net", "laravel", "gatsby", "remix",
    # Data / ML / AI
    "machine learning", "deep learning", "nlp", "computer vision", "tensorflow", "pytorch", "keras",
    "scikit-learn", "pandas", "numpy", "matplotlib", "seaborn", "plotly", "opencv", "huggingface",
    "langchain", "openai", "transformers", "xgboost", "lightgbm", "random forest", "neural network",
    "data analysis", "data science", "feature engineering", "model deployment",
    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "sqlite", "cassandra", "oracle", "dynamodb",
    "firebase", "elasticsearch", "neo4j", "influxdb", "supabase",
    # Cloud / DevOps
    "aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible", "jenkins", "ci/cd", "github actions",
    "linux", "nginx", "apache", "heroku", "vercel", "netlify", "cloudflare",
    # Tools
    "git", "github", "gitlab", "bitbucket", "jira", "confluence", "figma", "postman", "swagger",
    "grafana", "prometheus", "airflow", "spark", "hadoop", "kafka", "rabbitmq", "celery",
    # Mobile
    "android", "ios", "react native", "flutter", "xamarin",
    # Other
    "rest api", "graphql", "microservices", "agile", "scrum", "devops", "mlops", "llm",
    "excel", "power bi", "tableau", "looker", "dbt",
}

# ══════════════════════════════════════════════════════════
#  TEXT EXTRACTION
# ══════════════════════════════════════════════════════════


def extract_text_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        st.warning(f"PDF read error: {e}")
    return text.strip()


def extract_text_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        doc = docx_lib.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        st.warning(f"DOCX read error: {e}")
    return text.strip()


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_docx(file_bytes)
    return ""

# ══════════════════════════════════════════════════════════
#  PARSERS
# ══════════════════════════════════════════════════════════


def parse_email(text: str) -> str:
    match = re.search(
        r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def parse_phone(text: str) -> str:
    patterns = [
        r"(?:\+91[\s\-]?)?[6-9]\d{9}",
        r"\+?[\d][\d\s\-\(\)]{8,15}\d",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(0).strip()
    return ""


def parse_name(text: str, doc) -> str:
    # Try spaCy PERSON first
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            if 2 <= len(name.split()) <= 4 and len(name) < 50:
                return name
    # Fallback: first non-empty line that looks like a name
    for line in text.split("\n")[:8]:
        line = line.strip()
        if (2 <= len(line.split()) <= 4 and
                line.replace(" ", "").replace(".", "").isalpha() and
                len(line) < 50 and line[0].isupper()):
            return line
    return ""


def parse_skills(text: str) -> list:
    text_lower = text.lower()
    found = set()
    for skill in SKILLS_DB:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found.add(skill.title())
    return sorted(found)


def parse_education(text: str) -> list:
    degrees = [
        "b.tech", "m.tech", "b.e", "m.e", "bsc", "msc", "b.sc", "m.sc", "bca", "mca",
        "bba", "mba", "phd", "ph.d", "bachelor", "master", "diploma", "10th", "12th",
        "b.com", "m.com", "be", "me", "b.a", "m.a", "llb", "mbbs", "engineering",
    ]
    lines = text.split("\n")
    edu_lines = []
    capture = False
    for line in lines:
        l = line.strip()
        if not l:
            continue
        l_lower = l.lower()
        if any(kw in l_lower for kw in ["education", "academic", "qualification", "schooling"]):
            capture = True
            continue
        if capture and any(kw in l_lower for kw in ["experience", "project", "skill", "certification", "work", "employment"]):
            capture = False
        if capture and len(l) > 3:
            edu_lines.append(l)
        elif any(deg in l_lower for deg in degrees) and len(l) < 200:
            if l not in edu_lines:
                edu_lines.append(l)
    # Clean duplicates keeping order
    seen, result = set(), []
    for e in edu_lines:
        if e not in seen:
            seen.add(e)
            result.append(e)
    return result[:6]


def parse_experience(text: str) -> list:
    lines = text.split("\n")
    exp_lines = []
    capture = False
    for line in lines:
        l = line.strip()
        if not l:
            continue
        l_lower = l.lower()
        if any(kw in l_lower for kw in ["experience", "employment", "work history", "career", "professional"]):
            capture = True
            continue
        if capture and any(kw in l_lower for kw in ["education", "skill", "project", "certification", "academic"]):
            capture = False
        if capture and len(l) > 3:
            exp_lines.append(l)
    # Also catch year patterns like "2020 - 2023" lines near job titles
    year_pat = re.compile(r'\b(19|20)\d{2}\b')
    for i, l in enumerate(lines):
        if year_pat.search(l) and len(l.strip()) < 120:
            if l.strip() not in exp_lines:
                exp_lines.append(l.strip())
    seen, result = set(), []
    for e in exp_lines:
        if e not in seen:
            seen.add(e)
            result.append(e)
    return result[:10]


def parse_linkedin(text: str) -> str:
    m = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    return "https://" + m.group(0) if m else ""


def parse_github(text: str) -> str:
    m = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
    return "https://" + m.group(0) if m else ""


def completion_score(data: dict) -> int:
    fields = ["name", "email", "phone", "skills", "education", "experience"]
    filled = sum(1 for f in fields if data.get(
        f) and data[f] != "" and data[f] != [])
    return int((filled / len(fields)) * 100)


def parse_resume(text: str, filename: str) -> dict:
    doc = nlp(text[:50000])  # limit for performance
    skills = parse_skills(text)
    education = parse_education(text)
    experience = parse_experience(text)
    data = {
        "id":          str(uuid.uuid4())[:8],
        "filename":    filename,
        "name":        parse_name(text, doc),
        "email":       parse_email(text),
        "phone":       parse_phone(text),
        "linkedin":    parse_linkedin(text),
        "github":      parse_github(text),
        "skills":      skills,
        "education":   education,
        "experience":  experience,
        "parsed_at":   datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "raw_text":    text[:3000],
    }
    data["score"] = completion_score(data)
    return data


# ══════════════════════════════════════════════════════════
#  DATABASE
# ══════════════════════════════════════════════════════════
DB_PATH = "resumes.db"


def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS resumes (
            id          TEXT PRIMARY KEY,
            filename    TEXT,
            name        TEXT,
            email       TEXT,
            phone       TEXT,
            linkedin    TEXT,
            github      TEXT,
            skills      TEXT,
            education   TEXT,
            experience  TEXT,
            score       INTEGER,
            parsed_at   TEXT,
            raw_text    TEXT
        )""")
    conn.commit()
    conn.close()


def save_resume(data: dict):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        INSERT OR REPLACE INTO resumes
        (id,filename,name,email,phone,linkedin,github,skills,education,experience,score,parsed_at,raw_text)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (
        data["id"], data["filename"], data["name"], data["email"], data["phone"],
        data["linkedin"], data["github"],
        json.dumps(data["skills"]),
        json.dumps(data["education"]),
        json.dumps(data["experience"]),
        data["score"], data["parsed_at"], data["raw_text"],
    ))
    conn.commit()
    conn.close()


def fetch_all_resumes() -> list:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cur = conn.cursor()
    cur.execute("SELECT * FROM resumes ORDER BY parsed_at DESC")
    rows = []
    for r in cur.fetchall():
        d = dict(r)
        d["skills"] = json.loads(d["skills"] or "[]")
        d["education"] = json.loads(d["education"] or "[]")
        d["experience"] = json.loads(d["experience"] or "[]")
        rows.append(d)
    conn.close()
    return rows


def delete_resume(resume_id: str):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("DELETE FROM resumes WHERE id=?", (resume_id,))
    conn.commit()
    conn.close()


def clear_all_resumes():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("DELETE FROM resumes")
    conn.commit()
    conn.close()


def get_stats() -> dict:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM resumes")
    total = cur.fetchone()[0]
    cur.execute("SELECT AVG(score) FROM resumes")
    avg_score = cur.fetchone()[0] or 0
    conn.close()
    return {"total": total, "avg_score": round(avg_score)}


def resumes_to_df(resumes: list) -> pd.DataFrame:
    rows = []
    for r in resumes:
        rows.append({
            "ID":         r["id"],
            "Name":       r["name"],
            "Email":      r["email"],
            "Phone":      r["phone"],
            "Skills":     ", ".join(r["skills"][:8]),
            "Education":  " | ".join(r["education"][:2]),
            "Experience": " | ".join(r["experience"][:2]),
            "Score %":    r["score"],
            "File":       r["filename"],
            "Parsed At":  r["parsed_at"],
        })
    return pd.DataFrame(rows)


# ══════════════════════════════════════════════════════════
#  INIT
# ══════════════════════════════════════════════════════════
init_db()

if "parsed_results" not in st.session_state:
    st.session_state.parsed_results = []

# ══════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div style='text-align:center;padding:1rem 0 0.5rem;'>
        <div class='sb-title'>📄 ResumeParser</div>
        <div style='color:#64748b;font-size:0.78rem;'>AI-Powered · spaCy + Regex</div>
    </div><hr>
    """, unsafe_allow_html=True)

    stats = get_stats()
    st.markdown(f"""
    <div style='display:flex;gap:8px;margin-bottom:1rem;'>
        <div class='metric-card' style='flex:1;'>
            <div class='metric-val' style='color:#00d4aa;font-size:1.5rem;'>{stats['total']}</div>
            <div class='metric-label'>Resumes</div>
        </div>
        <div class='metric-card' style='flex:1;'>
            <div class='metric-val' style='color:#6c8dff;font-size:1.5rem;'>{stats['avg_score']}%</div>
            <div class='metric-label'>Avg Score</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("**📋 Extracts**")
    for item in ["✅ Name & Contact", "✅ Email & Phone",
                 "✅ LinkedIn & GitHub", "✅ Skills (50+ tech)",
                 "✅ Education", "✅ Work Experience"]:
        st.markdown(f"<div style='font-size:0.83rem;color:#94a3b8;padding:2px 0;'>{item}</div>",
                    unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    st.markdown("**📁 Formats Supported**")
    st.markdown("""
    <div style='font-size:0.83rem;color:#94a3b8;'>
        📕 PDF &nbsp; · &nbsp; 📘 DOCX &nbsp; · &nbsp; 📄 DOC
    </div>""", unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)
    if st.button("🗑️ Clear All Resumes", use_container_width=True):
        clear_all_resumes()
        st.session_state.parsed_results = []
        st.success("All records cleared!")
        st.rerun()

# ══════════════════════════════════════════════════════════
#  MAIN TABS
# ══════════════════════════════════════════════════════════
tab_upload, tab_search, tab_database, tab_guide = st.tabs([
    "📤  Upload & Parse",
    "🔍  Search & Filter",
    "🗄️  Database",
    "📖  How to Use",
])

# ══════════════════════════════════════════════════════════
#  TAB 1 — UPLOAD & PARSE
# ══════════════════════════════════════════════════════════
with tab_upload:
    st.markdown("""
    <div class="hero">
        <div class="hero-title">📄 Resume Parser AI</div>
        <div class="hero-sub">Upload resumes · Extract info instantly · Store & export candidates</div>
    </div>
    """, unsafe_allow_html=True)

    uploaded_files = st.file_uploader(
        "Drop your resume files here",
        type=["pdf", "docx", "doc"],
        accept_multiple_files=True,
        help="Supports PDF, DOCX, DOC — upload multiple files at once!",
    )

    if uploaded_files:
        col_parse, col_clear = st.columns([1, 5])
        with col_parse:
            parse_btn = st.button(
                "⚡ Parse Resumes", type="primary", use_container_width=True)

        if parse_btn:
            st.session_state.parsed_results = []
            progress = st.progress(0, text="Parsing resumes…")
            for i, uf in enumerate(uploaded_files):
                progress.progress((i + 1) / len(uploaded_files),
                                  text=f"Parsing {uf.name}… ({i+1}/{len(uploaded_files)})")
                file_bytes = uf.read()
                text = extract_text(file_bytes, uf.name)
                if text:
                    result = parse_resume(text, uf.name)
                    save_resume(result)
                    st.session_state.parsed_results.append(result)
                else:
                    st.warning(f"⚠️ Could not extract text from `{uf.name}`")
            progress.empty()
            st.success(
                f"✅ Parsed {len(st.session_state.parsed_results)} resume(s) successfully!")
            st.rerun()

    # ── Display Results ──────────────────────────────────
    if st.session_state.parsed_results:
        st.markdown("<hr>", unsafe_allow_html=True)
        st.markdown(f"""
        <div class='sec-head'>
            Parsed Results
            <span class='badge-new'>{len(st.session_state.parsed_results)} resume(s)</span>
        </div>
        """, unsafe_allow_html=True)

        for idx, res in enumerate(st.session_state.parsed_results):
            score = res["score"]
            score_color = "#00d4aa" if score >= 70 else "#f59e0b" if score >= 40 else "#f87171"

            with st.expander(
                f"📋 {res['name'] or res['filename']}  —  Score: {score}%",
                expanded=(idx == 0)
            ):
                # Score bar
                st.markdown(f"""
                <div class='parse-score'>
                    <span style='font-size:0.8rem;color:#64748b;'>Completeness</span>
                    <div class='score-bar-wrap'>
                        <div class='score-bar-fill' style='width:{score}%;background:linear-gradient(90deg,{score_color},{score_color}99);'></div>
                    </div>
                    <span class='score-pct' style='color:{score_color};'>{score}%</span>
                </div>
                """, unsafe_allow_html=True)

                c1, c2 = st.columns(2)

                with c1:
                    # Contact Info Card
                    st.markdown("""<div class='card'><div class='card-title'>👤 Contact Information</div>""",
                                unsafe_allow_html=True)
                    info_rows = [
                        ("Name",     res["name"] or "—"),
                        ("Email",    res["email"] or "—"),
                        ("Phone",    res["phone"] or "—"),
                        ("LinkedIn",
                         f'<a href="{res["linkedin"]}" target="_blank" style="color:#6c8dff;">{res["linkedin"]}</a>' if res["linkedin"] else "—"),
                        ("GitHub",
                         f'<a href="{res["github"]}" target="_blank" style="color:#6c8dff;">{res["github"]}</a>' if res["github"] else "—"),
                        ("File",     res["filename"]),
                        ("Parsed",   res["parsed_at"]),
                    ]
                    for label, val in info_rows:
                        st.markdown(f"""
                        <div class='info-row'>
                            <span class='info-label'>{label}</span>
                            <span class='info-val'>{val}</span>
                        </div>
                        """, unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)

                    # Education Card
                    st.markdown("""<div class='card'><div class='card-title'>🎓 Education</div>""",
                                unsafe_allow_html=True)
                    if res["education"]:
                        for edu in res["education"]:
                            st.markdown(f"""
                            <div class='info-row'>
                                <span class='info-val'>• {edu}</span>
                            </div>""", unsafe_allow_html=True)
                    else:
                        st.markdown("<div style='color:#64748b;font-size:0.85rem;'>Not detected</div>",
                                    unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)

                with c2:
                    # Skills Card
                    st.markdown("""<div class='card'><div class='card-title'>⚡ Skills & Technologies</div>""",
                                unsafe_allow_html=True)
                    if res["skills"]:
                        tags_html = "".join(
                            f"<span class='skill-tag'>{s}</span>" for s in res["skills"])
                        st.markdown(
                            f"<div class='skill-wrap'>{tags_html}</div>", unsafe_allow_html=True)
                    else:
                        st.markdown("<div style='color:#64748b;font-size:0.85rem;'>No skills detected</div>",
                                    unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)

                    # Experience Card
                    st.markdown("""<div class='card'><div class='card-title'>💼 Work Experience</div>""",
                                unsafe_allow_html=True)
                    if res["experience"]:
                        for exp in res["experience"][:8]:
                            st.markdown(f"""
                            <div class='info-row'>
                                <span class='info-val'>• {exp}</span>
                            </div>""", unsafe_allow_html=True)
                    else:
                        st.markdown("<div style='color:#64748b;font-size:0.85rem;'>Not detected</div>",
                                    unsafe_allow_html=True)
                    st.markdown("</div>", unsafe_allow_html=True)

                # Download this resume's JSON
                st.download_button(
                    "⬇️ Download as JSON",
                    data=json.dumps(
                        {k: v for k, v in res.items() if k != "raw_text"}, indent=2),
                    file_name=f"{res['name'] or res['id']}_parsed.json",
                    mime="application/json",
                    key=f"json_{res['id']}",
                )

# ══════════════════════════════════════════════════════════
#  TAB 2 — SEARCH & FILTER
# ══════════════════════════════════════════════════════════
with tab_search:
    st.markdown("<div class='sec-head'>🔍 Search & Filter Candidates</div>",
                unsafe_allow_html=True)

    all_resumes = fetch_all_resumes()

    if not all_resumes:
        st.info(
            "💡 No resumes in the database yet. Upload some in the **Upload & Parse** tab!")
    else:
        # ── Filters ─────────────────────────────────────
        with st.container():
            f1, f2, f3 = st.columns([2, 2, 1])
            with f1:
                search_name = st.text_input(
                    "🔎 Search by Name / Email", placeholder="e.g. John, john@email.com")
            with f2:
                search_skill = st.text_input(
                    "⚡ Filter by Skill", placeholder="e.g. Python, React, SQL")
            with f3:
                min_score = st.slider("Min Score %", 0, 100, 0, 10)

        # ── Apply Filters ────────────────────────────────
        filtered = all_resumes
        if search_name:
            q = search_name.lower()
            filtered = [r for r in filtered if
                        q in r["name"].lower() or q in r["email"].lower()]
        if search_skill:
            q = search_skill.lower()
            filtered = [r for r in filtered if
                        any(q in s.lower() for s in r["skills"])]
        if min_score > 0:
            filtered = [r for r in filtered if r["score"] >= min_score]

        st.markdown(f"<div style='color:#64748b;font-size:0.85rem;margin-bottom:1rem;'>"
                    f"Showing <strong style='color:#00d4aa;'>{len(filtered)}</strong> of "
                    f"<strong>{len(all_resumes)}</strong> candidates</div>",
                    unsafe_allow_html=True)

        if not filtered:
            st.warning("No candidates match your filters.")
        else:
            # Summary table
            df_filtered = resumes_to_df(filtered)
            st.dataframe(df_filtered, use_container_width=True,
                         hide_index=True, height=350)

            # CSV Export of filtered results
            csv = df_filtered.to_csv(index=False).encode("utf-8")
            st.download_button(
                "⬇️ Download Filtered Results as CSV",
                data=csv,
                file_name=f"candidates_filtered_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
            )

            st.markdown("<hr>", unsafe_allow_html=True)

            # Candidate cards
            st.markdown("**Candidate Cards**")
            for r in filtered:
                score = r["score"]
                score_color = "#00d4aa" if score >= 70 else "#f59e0b" if score >= 40 else "#f87171"
                with st.expander(f"👤 {r['name'] or 'Unknown'}  |  {r['email'] or '—'}  |  Score: {score}%"):
                    ca, cb = st.columns(2)
                    with ca:
                        st.markdown(f"**📞 Phone:** {r['phone'] or '—'}")
                        st.markdown(f"**📅 Parsed:** {r['parsed_at']}")
                        st.markdown(f"**📁 File:** {r['filename']}")
                        if r["linkedin"]:
                            st.markdown(
                                f"**🔗 LinkedIn:** [{r['linkedin']}]({r['linkedin']})")
                        if r["education"]:
                            st.markdown("**🎓 Education:**")
                            for e in r["education"][:3]:
                                st.markdown(f"  - {e}")
                    with cb:
                        st.markdown(f"**⚡ Skills ({len(r['skills'])}):**")
                        tags = "".join(
                            f"<span class='skill-tag'>{s}</span>" for s in r["skills"][:15])
                        st.markdown(
                            f"<div class='skill-wrap'>{tags}</div>", unsafe_allow_html=True)
                    del_col, _ = st.columns([1, 5])
                    with del_col:
                        if st.button("🗑️ Delete", key=f"del_{r['id']}"):
                            delete_resume(r["id"])
                            st.rerun()

# ══════════════════════════════════════════════════════════
#  TAB 3 — DATABASE
# ══════════════════════════════════════════════════════════
with tab_database:
    st.markdown("<div class='sec-head'>🗄️ Candidate Database</div>",
                unsafe_allow_html=True)

    all_resumes = fetch_all_resumes()
    stats = get_stats()

    # Metric cards
    mc1, mc2, mc3, mc4 = st.columns(4)
    all_skills_flat = [s for r in all_resumes for s in r["skills"]]
    top_skill = max(set(all_skills_flat),
                    key=all_skills_flat.count) if all_skills_flat else "—"
    avg_skills = round(sum(len(r["skills"])
                       for r in all_resumes) / max(len(all_resumes), 1))

    for col, val, color, label in [
        (mc1, stats["total"],      "#00d4aa", "Total Candidates"),
        (mc2, f"{stats['avg_score']}%", "#6c8dff", "Avg Completeness"),
        (mc3, avg_skills,          "#f59e0b", "Avg Skills/Resume"),
        (mc4, top_skill,           "#f472b6", "Top Skill"),
    ]:
        with col:
            st.markdown(f"""
            <div class='metric-card'>
                <div class='metric-val' style='color:{color};font-size:1.6rem;'>{val}</div>
                <div class='metric-label'>{label}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("<hr>", unsafe_allow_html=True)

    if not all_resumes:
        st.info("💡 No data yet. Upload resumes in the **Upload & Parse** tab!")
    else:
        df_all = resumes_to_df(all_resumes)
        st.dataframe(df_all, use_container_width=True,
                     height=400, hide_index=True)

        dl_c1, dl_c2 = st.columns(2)
        with dl_c1:
            csv_all = df_all.to_csv(index=False).encode("utf-8")
            st.download_button(
                "⬇️ Download All as CSV",
                data=csv_all,
                file_name=f"all_candidates_{datetime.now().strftime('%Y%m%d_%H%M')}.csv",
                mime="text/csv",
                use_container_width=True,
            )
        with dl_c2:
            json_all = json.dumps(
                [{k: v for k, v in r.items() if k != "raw_text"}
                 for r in all_resumes],
                indent=2
            ).encode("utf-8")
            st.download_button(
                "⬇️ Download All as JSON",
                data=json_all,
                file_name=f"all_candidates_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True,
            )

        # Skills frequency chart
        if all_skills_flat:
            st.markdown("<hr>", unsafe_allow_html=True)
            st.markdown("**⚡ Top Skills Across All Candidates**")
            from collections import Counter
            skill_counts = Counter(all_skills_flat).most_common(15)
            df_skills = pd.DataFrame(skill_counts, columns=["Skill", "Count"])
            st.bar_chart(df_skills.set_index("Skill"),
                         color="#00d4aa", use_container_width=True)

# ══════════════════════════════════════════════════════════
#  TAB 4 — HOW TO USE
# ══════════════════════════════════════════════════════════
with tab_guide:
    st.markdown("<div class='sec-head'>📖 How to Use Resume Parser AI</div>",
                unsafe_allow_html=True)

    st.markdown("""
    <div class='card'>
        <div class='card-title'>🚀 Quick Start</div>
        <div class='info-row'><span class='info-label'>Step 1</span>
            <span class='info-val'>Go to the <strong>Upload & Parse</strong> tab</span></div>
        <div class='info-row'><span class='info-label'>Step 2</span>
            <span class='info-val'>Drag & drop one or multiple PDF/DOCX resumes</span></div>
        <div class='info-row'><span class='info-label'>Step 3</span>
            <span class='info-val'>Click <strong>⚡ Parse Resumes</strong> — extraction happens instantly</span></div>
        <div class='info-row'><span class='info-label'>Step 4</span>
            <span class='info-val'>View extracted data, then search/filter in the <strong>Search & Filter</strong> tab</span></div>
        <div class='info-row'><span class='info-label'>Step 5</span>
            <span class='info-val'>Download results as <strong>CSV or JSON</strong> from the Database tab</span></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class='card'>
        <div class='card-title'>🧠 How Extraction Works</div>
        <div class='info-row'><span class='info-label'>Name</span>
            <span class='info-val'>spaCy NER (PERSON entity) + first-line heuristic fallback</span></div>
        <div class='info-row'><span class='info-label'>Email</span>
            <span class='info-val'>Regex pattern matching standard email formats</span></div>
        <div class='info-row'><span class='info-label'>Phone</span>
            <span class='info-val'>Regex with Indian & international number patterns</span></div>
        <div class='info-row'><span class='info-label'>Skills</span>
            <span class='info-val'>Keyword matching against 150+ tech skills dictionary</span></div>
        <div class='info-row'><span class='info-label'>Education</span>
            <span class='info-val'>Section detection + degree keyword matching</span></div>
        <div class='info-row'><span class='info-label'>Experience</span>
            <span class='info-val'>Section detection + year pattern recognition</span></div>
        <div class='info-row'><span class='info-label'>Score</span>
            <span class='info-val'>Completeness % — how many of 6 fields were successfully extracted</span></div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class='card'>
        <div class='card-title'>💡 Tips for Best Results</div>
        <div class='info-row'><span class='info-val'>✅ Use text-based PDFs (not scanned images)</span></div>
        <div class='info-row'><span class='info-val'>✅ Resumes with clear section headers work best (Education, Experience, Skills)</span></div>
        <div class='info-row'><span class='info-val'>✅ Upload multiple resumes at once for batch processing</span></div>
        <div class='info-row'><span class='info-val'>✅ Use the Search tab to filter by skill for quick shortlisting</span></div>
        <div class='info-row'><span class='info-val'>⚠️ Scanned/image PDFs won't extract text (OCR not included in this version)</span></div>
    </div>
    """, unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════
#  Footer
# ══════════════════════════════════════════════════════════
st.markdown("""
<div style='text-align:center;margin-top:3rem;color:#1e293b;font-size:0.78rem;'>
    Built with 🐍 Python · spaCy · pdfplumber · Streamlit · SQLite &nbsp;|&nbsp; ResumeParser AI v1.0
</div>
""", unsafe_allow_html=True)
